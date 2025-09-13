if __name__ == "__main__":
    """
        生成三元组
    """
    from openai import OpenAI
    import os
    import json
    import glob
    import re
    from datetime import datetime
    from pdf2image import convert_from_path
    from docx import Document
    import openpyxl
    from PIL import ImageGrab
    import io


    # 查找最新的 data_时间戳.json 文件
    def find_latest_data_file(directory="received_jsons"):
        # 获取所有 data_*.json 文件
        pattern = os.path.join(directory, "data_*.json")
        data_files = glob.glob(pattern)
        
        if not data_files:
            raise FileNotFoundError(f"在 {directory} 目录中未找到 data_*.json 文件")
        
        # 从文件名中提取时间戳并找到最新的文件
        def extract_timestamp(filename):
            # 从文件名中提取数字部分作为时间戳
            match = re.search(r'data_(\d+)\.json', os.path.basename(filename))
            if match:
                return int(match.group(1))
            return 0
        
        # 按时间戳排序，获取最新的文件
        latest_file = max(data_files, key=extract_timestamp)
        return latest_file

    # 从最新的 data_*.json 文件中提取信息
    try:
        latest_data_file = find_latest_data_file()
        print(f"找到最新的数据文件: {latest_data_file}")
        
        with open(latest_data_file, "r", encoding="utf-8") as f:
            data_content = json.load(f)
        
        # 根据您提供的示例，JSON结构是直接包含字段，而不是嵌套在data键中
        name = data_content.get("name", "未知文件")
        file_name = data_content.get("fileName", "unknown")
        file_type = data_content.get("fileType", "txt")
        
        print(f"提取的信息 - name: {name}, fileName: {file_name}, fileType: {file_type}")
        
    except Exception as e:
        print(f"错误: {e}")
        # 如果找不到文件或解析失败，使用默认值
        name = "未知文件"
        file_name = "unknown"
        file_type = "txt"

    import os
    import glob
    from docx import Document
    import openpyxl

    def extract_text_from_file(file_path):
        """从文件中提取文本内容"""
        text = ""
        
        if file_path.endswith('.docx'):
            # 提取Word文本
            try:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():  # 忽略空行
                        text += para.text + "\n"
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            text += " | ".join(row_text) + "\n"
            except Exception as e:
                print(f"处理Word文件 {file_path} 时出错: {e}")
        
        elif file_path.endswith('.xlsx'):
            # 提取Excel文本
            try:
                wb = openpyxl.load_workbook(file_path)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text += f"\n--- 工作表: {sheet_name} ---\n"
                    for row in ws.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None and str(cell).strip():
                                row_text.append(str(cell))
                        if row_text:
                            text += " | ".join(row_text) + "\n"
            except Exception as e:
                print(f"处理Excel文件 {file_path} 时出错: {e}")
        
        elif file_path.endswith('.txt'):
            # 提取纯文本文件内容
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # 如果UTF-8解码失败，尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        text = f.read()
                except Exception as e:
                    print(f"处理文本文件 {file_path} 时出错: {e}")
            except Exception as e:
                print(f"处理文本文件 {file_path} 时出错: {e}")
        
        elif file_path.endswith('.md'):
            # 提取Markdown文件内容
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # 如果UTF-8解码失败，尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        text = f.read()
                except Exception as e:
                    print(f"处理Markdown文件 {file_path} 时出错: {e}")
            except Exception as e:
                print(f"处理Markdown文件 {file_path} 时出错: {e}")
        
        return text

    import os
    import glob

    # 获取input目录下所有支持的文件（包括子目录）
    input_dir = "input"

    # 使用递归通配符模式匹配所有子目录中的文件
    excel_files = glob.glob(os.path.join(input_dir, "**", "*.xlsx"), recursive=True)
    word_files = glob.glob(os.path.join(input_dir, "**", "*.docx"), recursive=True)
    text_files = glob.glob(os.path.join(input_dir, "**", "*.txt"), recursive=True)
    markdown_files = glob.glob(os.path.join(input_dir, "**", "*.md"), recursive=True)

    all_files = excel_files + word_files + text_files + markdown_files

    # 提取所有文件的文本内容
    all_text = ""
    for file_path in all_files:
        # 获取相对于input目录的路径，以保留子目录结构信息
        relative_path = os.path.relpath(file_path, input_dir)
        print(f"正在处理: {relative_path}")
        file_text = extract_text_from_file(file_path)
        all_text += f"\n\n--- 文件: {relative_path} ---\n\n"
        all_text += file_text


    client = OpenAI(
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        api_key="sk-94008d6ff04d42e583e4aa661694b92a",  # 替换为你的实际API Key
    )

    completion = client.chat.completions.create(
        model="qwen-vl-plus",  # 此处以qwen-vl-plus为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=[
            {
                "role": "system",
                "content": (
                    "你是一个信息提取专家，需要将用户提供的文本转换成三元组JSON格式。"
                    "输出格式要求："
                    "1. 只输出'triples'数组部分"
                    "2. 'triples'的值是一个数组，每个元素是一个对象"
                    "3. 每个对象包含三个键: 'subject'(主语), 'predicate'(关系), 'object'(宾语)"
                    "4. 每个键值对必须单独一行显示"
                    "5. 键和值之间保留一个空格"
                    "6. 只输出核心事实信息，忽略次要细节"
                    "7. 确保输出是有效的JSON格式"
                    "8. 输出格式必须严格遵循以下示例：\n"
                    "[\n"
                    "  {\n"
                    "   \"subject\": \"北京\", \n"
                    "   \"predicate\": \"是...的首都\", \n"
                    "   \"object\": \"中国\"\n"
                    "  },\n"
                    "  {\n"
                    "   \"subject\": \"李白\", \n"
                    "   \"predicate\": \"是...的诗人\", \n"
                    "   \"object\": \"唐朝\"\n"
                    "  }\n"
                    "]"
                )
            },
            {
                "role": "user",
                "content": (
                    f"请将以下文本转换为三元组JSON格式：\n\n{all_text}\n\n"
                    "严格按照上述示例的格式输出，只返回'triples'数组部分。"
                )
            }
        ],
        temperature=0.1,  # 更低的温度确保格式精确
        response_format={"type": "json_object"}  # 强制JSON输出
    )

    # 提取模型返回的triples数组
    triples_result = completion.choices[0].message.content

    # 构建完整的输出结构
    full_output = {
        "code": "104",
        "message": "三元组数据",
        "data": {
            "name": name,  # 从最新的 data_*.json 文件中提取
            "fileName": file_name,  # 从最新的 data_*.json 文件中提取
            "fileType": file_type,  # 从最新的 data_*.json 文件中提取
            "triples": []  # 将由模型结果填充
        }
    }

    # 将模型返回的triples字符串转换为Python对象
    try:
        # 解析模型返回的JSON
        model_output = json.loads(triples_result)
        
        # 如果模型返回的是包含"triples"键的对象
        if isinstance(model_output, dict) and "triples" in model_output:
            full_output["data"]["triples"] = model_output["triples"]
        # 如果模型返回的是数组（直接triples数组）
        elif isinstance(model_output, list):
            full_output["data"]["triples"] = model_output
        else:
            print("警告：模型返回的格式不符合预期，将使用空数组")
            full_output["data"]["triples"] = []
    except Exception as e:
        print(f"解析模型输出时出错: {e}")
        full_output["data"]["triples"] = []

    # 生成输出文件名，使用当前时间戳
    output_filename = "demo.json"

    # 将完整结果保存为文件
    with open(output_filename, "w", encoding="utf-8") as out_file:
        out_file.write('{\n')
        out_file.write('    "code": "104",\n')
        out_file.write('    "message": "三元组数据",\n')
        out_file.write('    "data": {\n')
        out_file.write(f'        "name": "{full_output["data"]["name"]}",\n')
        out_file.write(f'        "fileName": "{full_output["data"]["fileName"]}",\n')
        out_file.write(f'        "fileType": "{full_output["data"]["fileType"]}",\n')
        out_file.write('        "triples": [\n')
        
        # 格式化triples数组
        for i, triple in enumerate(full_output["data"]["triples"]):
            out_file.write('            {\n')
            out_file.write(f'                "subject": "{triple["subject"]}",\n')
            out_file.write(f'                "predicate": "{triple["predicate"]}",\n')
            # 最后一个元素不加逗号
            if i < len(full_output["data"]["triples"]) - 1:
                out_file.write(f'                "object": "{triple["object"]}"\n')
                out_file.write('            },\n')
            else:
                out_file.write(f'                "object": "{triple["object"]}"\n')
                out_file.write('            }\n')
        
        out_file.write('        ]\n')
        out_file.write('    }\n')
        out_file.write('}')

    print(f"三元组已成功保存到 {output_filename} 文件，格式完全符合要求")